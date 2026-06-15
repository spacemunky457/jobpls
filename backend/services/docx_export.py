"""Render a plain-text tailored CV into a .docx.

Mirrors the layout heuristics of services/pdf.py: the first non-blank line is
the name, an immediately following line is a tagline, ALL-CAPS short lines
become section headings. Unlike the PDF path, .docx handles Unicode natively
so no transliteration is needed."""

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from services.pdf import _is_heading


def cv_text_to_docx(text: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.6)
        section.left_margin = section.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(2)

    lines = (text or "").split("\n")
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        run = doc.add_paragraph().add_run(lines[i].strip())
        run.bold = True
        run.font.size = Pt(15)
        i += 1
        # A following non-blank line is usually a title/tagline.
        if i < len(lines) and lines[i].strip():
            run = doc.add_paragraph().add_run(lines[i].strip())
            run.font.color.rgb = RGBColor(90, 90, 90)
            i += 1

    for line in lines[i:]:
        stripped = line.rstrip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if _is_heading(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
        else:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
